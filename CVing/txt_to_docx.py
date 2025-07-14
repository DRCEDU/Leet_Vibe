#!/usr/bin/env python3
"""
Text to Word Document Converter

This script converts a text file to a formatted Word document (.docx).
It includes basic formatting for CV/resume documents.

Usage:
    python txt_to_docx.py <input_text_file> [output_docx_file]
"""

import sys
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("Error: python-docx library not found.")
    print("Please install it using: pip install python-docx")
    sys.exit(1)


def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add color (blue) and underline
    c = OxmlElement('w:color')
    c.set(qn('w:val'), "0563C1")
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def format_cv_document(text_content, output_path):
    """
    Convert text content to a formatted Word document.
    """
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    lines = text_content.split('\n')
    current_section = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Main heading (Name)
        if line == "C.J. Duan":
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(18)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Contact information
        elif "Email:" in line and "LinkedIn:" in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Parse contact info
            parts = line.split(" | ")
            for i, part in enumerate(parts):
                if i > 0:
                    p.add_run(" | ")
                
                if "Email:" in part:
                    p.add_run("Email: ")
                    email = part.split("Email: ")[1]
                    add_hyperlink(p, f"mailto:{email}", email)
                elif "LinkedIn:" in part:
                    p.add_run("LinkedIn: ")
                    linkedin_url = part.split("LinkedIn: ")[1]
                    add_hyperlink(p, f"https://{linkedin_url}", linkedin_url)
                elif "Website:" in part:
                    p.add_run("Website: ")
                    website = part.split("Website: ")[1]
                    add_hyperlink(p, f"https://{website}", website)
                else:
                    p.add_run(part)
            
            doc.add_paragraph()  # Add space after contact info
            
        # Section headers
        elif line in ["Professional Summary", "Technical Skills and Toolsets", "Professional Experience", 
                     "Project Highlights", "Education", "Selected Publications & Research", "Contact & References"]:
            current_section = line
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(14)
            run.bold = True
            run.underline = True
            
        # Job titles/positions
        elif any(title in line for title in ["Professor", "Scientist", "Research:"]):
            if not line.startswith("-"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(12)
                
        # Company names and dates
        elif any(org in line for org in ["University", "PepsiCo", "Troy University", "Clemson University", "DRC Lab"]):
            if not line.startswith("-") and "Professor" not in line:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.italic = True
                
        # Date ranges
        elif any(month in line for month in ["January", "February", "March", "April", "May", "June",
                                           "July", "August", "September", "October", "November", "December"]) and "–" in line:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.italic = True
            
        # Bullet points
        elif line.startswith("-"):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            
        # Project names
        elif line.startswith("Project:"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            
        # Regular paragraphs
        else:
            if current_section == "Professional Summary":
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                p = doc.add_paragraph(line)
    
    # Save the document
    doc.save(output_path)
    print(f"Formatted CV saved to: {output_path}")


def main():
    """Main function to handle command line arguments and convert the file."""
    
    if len(sys.argv) < 2:
        print("Usage: python txt_to_docx.py <input_text_file> [output_docx_file]")
        print("\nExample:")
        print("python txt_to_docx.py cv_extracted.txt cv_formatted.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    # Generate output filename if not provided
    if len(sys.argv) > 2:
        output_file = sys.argv[2]
    else:
        input_path = Path(input_file)
        output_file = input_path.parent / f"{input_path.stem}_formatted.docx"
    
    # Check if input file exists
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found.")
        sys.exit(1)
    
    # Read the text content
    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            text_content = f.read()
    except Exception as e:
        print(f"Error reading input file: {str(e)}")
        sys.exit(1)
    
    # Convert to Word document
    format_cv_document(text_content, output_file)


if __name__ == "__main__":
    main()
