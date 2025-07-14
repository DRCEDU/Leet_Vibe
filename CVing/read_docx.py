#!/usr/bin/env python3
"""
Word Document Reader Script

This script reads .docx files and extracts their text content.
It can handle paragraphs, tables, and basic formatting.

Usage:
    python read_docx.py <path_to_docx_file>
"""

import sys
import os
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Error: python-docx library not found.")
    print("Please install it using: pip install python-docx")
    sys.exit(1)


def read_docx_file(file_path):
    """
    Read a .docx file and extract all text content.
    
    Args:
        file_path (str): Path to the .docx file
        
    Returns:
        str: Extracted text content
    """
    try:
        # Load the document
        doc = Document(file_path)
        
        # Extract text from paragraphs
        text_content = []
        
        print(f"Reading document: {file_path}")
        print("=" * 50)
        
        # Read paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text_content.append(paragraph.text)
        
        # Read tables if any
        if doc.tables:
            text_content.append("\n--- TABLES ---")
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_content.append(" | ".join(row_text))
        
        return "\n".join(text_content)
        
    except Exception as e:
        return f"Error reading document: {str(e)}"


def save_text_to_file(text_content, original_file_path):
    """
    Save extracted text to a .txt file.
    
    Args:
        text_content (str): The extracted text
        original_file_path (str): Original .docx file path
    """
    # Create output filename
    original_path = Path(original_file_path)
    output_file = original_path.parent / f"{original_path.stem}_extracted.txt"
    
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text_content)
        print(f"\nText content saved to: {output_file}")
    except Exception as e:
        print(f"Error saving file: {str(e)}")


def main():
    """Main function to handle command line arguments and process the document."""
    
    if len(sys.argv) != 2:
        print("Usage: python read_docx.py <path_to_docx_file>")
        print("\nExample:")
        print("python read_docx.py /path/to/document.docx")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    # Check if file exists
    if not os.path.exists(file_path):
        print(f"Error: File '{file_path}' not found.")
        sys.exit(1)
    
    # Check if file has .docx extension
    if not file_path.lower().endswith('.docx'):
        print("Warning: File doesn't have .docx extension. Proceeding anyway...")
    
    # Read the document
    text_content = read_docx_file(file_path)
    
    # Display the content
    print(text_content)
    
    # Ask if user wants to save to file
    while True:
        save_choice = input("\nSave extracted text to file? (y/n): ").lower().strip()
        if save_choice in ['y', 'yes']:
            save_text_to_file(text_content, file_path)
            break
        elif save_choice in ['n', 'no']:
            print("Text not saved.")
            break
        else:
            print("Please enter 'y' or 'n'")


if __name__ == "__main__":
    main()
